"""Phase 7 — Document processing & chunking tests."""
from __future__ import annotations

import asyncio
import io
import os
import uuid
from typing import AsyncIterator

import pytest
import pytest_asyncio
from fastapi.testclient import TestClient

from app.middleware.org_context import OrgContext
from app.services.document_processing import (
    CHUNK_OVERLAP_CHARS,
    ExtractedPage,
    TARGET_CHUNK_CHARS,
    chunk_pages,
    extract_text,
    normalize,
)


# ─────────────────── Unit (no DB) ───────────────────

def test_extract_txt_roundtrip():
    pages = extract_text("notes.txt", "text/plain", b"hello world")
    assert len(pages) == 1
    assert pages[0].text == "hello world"
    assert pages[0].page == 1


def test_extract_md_picks_section_heading():
    md = b"# Onboarding\n\nWelcome to OraOne.\n\n## Step 1\n\nDo a thing."
    pages = extract_text("welcome.md", "text/markdown", md)
    assert pages[0].section == "Onboarding"


def test_extract_csv_flattens_rows():
    csv_data = b"name,role\nAlice,owner\nBob,member\n"
    pages = extract_text("team.csv", "text/csv", csv_data)
    assert "name: Alice" in pages[0].text
    assert "role: member" in pages[0].text


def test_extract_pdf_multi_page():
    """Generate a tiny 3-page PDF with pypdf itself, then extract it."""
    from pypdf import PdfReader, PdfWriter
    from pypdf.generic import RectangleObject

    writer = PdfWriter()
    for i in range(3):
        writer.add_blank_page(width=200, height=200)
    buf = io.BytesIO()
    writer.write(buf)
    pdf_bytes = buf.getvalue()

    pages = extract_text("blank.pdf", "application/pdf", pdf_bytes)
    assert len(pages) == 3
    assert [p.page for p in pages] == [1, 2, 3]
    # pypdf just confirmed we got per-page pages back; content is empty
    # because the pages were blank by design — the structural assertion
    # is what matters here.
    assert RectangleObject  # silence unused-import lint


def test_extract_docx_with_heading():
    """Build a DOCX in memory and extract its sectioned paragraphs."""
    from docx import Document as DocxDocument

    docx_doc = DocxDocument()
    docx_doc.add_heading("Refund Policy", level=1)
    docx_doc.add_paragraph("All refunds within 30 days.")
    docx_doc.add_paragraph("Exceptions apply.")
    buf = io.BytesIO()
    docx_doc.save(buf)

    pages = extract_text("policy.docx", None, buf.getvalue())
    assert len(pages) == 1
    assert "Refund Policy" in pages[0].text
    assert pages[0].section == "Refund Policy"


def test_extract_rejects_unknown_type():
    with pytest.raises(ValueError):
        extract_text("data.bin", "application/octet-stream", b"\x00\x01")


def test_normalize_collapses_whitespace():
    assert normalize("hello\r\nworld\n\n\n\nbye  \tnow") == "hello\nworld\n\nbye now"


def test_chunk_respects_target_and_overlap():
    long_para = "Lorem ipsum dolor sit amet. " * 200  # ~5600 chars
    chunks = chunk_pages(
        [ExtractedPage(page=1, text=long_para)],
        source_file="x.txt",
        target_chars=400,
        overlap_chars=80,
    )

    assert len(chunks) >= 4
    # No chunk exceeds target_chars
    for c in chunks:
        assert len(c.content) <= 400

    # Consecutive chunks must overlap — the last 80 chars of chunk[i]
    # should appear inside chunk[i+1]. Use a soft check (any 30-char
    # window match) to tolerate boundary snapping.
    for i in range(len(chunks) - 1):
        tail = chunks[i].content[-80:]
        # Find any 30-char window from the tail inside the next chunk
        windows = [tail[j : j + 30] for j in range(0, max(1, len(tail) - 30))]
        assert any(w and w in chunks[i + 1].content for w in windows), (
            f"chunks {i} and {i+1} have no overlapping text — "
            f"chunk {i} ends with {tail!r}"
        )


def test_chunk_metadata_carries_required_keys():
    pages = [
        ExtractedPage(page=2, text="# Intro\n\nThis is page two.", section="Top"),
    ]
    chunks = chunk_pages(pages, source_file="report.pdf")
    assert chunks, "expected at least one chunk"
    md = chunks[0].metadata
    assert md["page"] == 2
    assert md["source_file"] == "report.pdf"
    # section is whichever heuristic kicks in first — must be non-empty
    assert md["section"]


def test_chunk_index_is_monotonic_and_zero_based():
    txt = "Sentence. " * 500  # ~5000 chars
    chunks = chunk_pages(
        [ExtractedPage(page=1, text=txt)],
        source_file="s.txt",
        target_chars=TARGET_CHUNK_CHARS,
        overlap_chars=CHUNK_OVERLAP_CHARS,
    )
    assert chunks[0].index == 0
    indices = [c.index for c in chunks]
    assert indices == sorted(indices)
    assert indices == list(range(len(chunks)))


def test_swagger_documents_process_endpoint():
    from server import app

    paths = app.openapi()["paths"]
    assert "/api/documents/{document_id}/process" in paths
    op = paths["/api/documents/{document_id}/process"]["post"]
    assert op["summary"]
    # DocumentRead now has the telemetry fields
    schema = app.openapi()["components"]["schemas"]["DocumentRead"]
    for f in ("processing_started_at", "processing_completed_at",
              "processing_error", "processing_time_ms"):
        assert f in schema["properties"], f"missing field: {f}"


# ─────────────────── DB-backed integration ───────────────────

def _ctx(org_id: uuid.UUID, *, user_id: uuid.UUID | None = None) -> OrgContext:
    return OrgContext(
        user_id=user_id or uuid.uuid4(),
        cognito_sub="sub-" + str(org_id),
        organization_id=org_id,
        membership_role="owner",
    )


def _postgres_reachable() -> bool:
    import asyncpg

    url = os.environ.get("DATABASE_URL", "")
    if not url:
        return False
    dsn = (
        url.replace("postgresql+asyncpg://", "postgresql://")
        .replace("postgresql+psycopg2://", "postgresql://")
    )

    async def _probe() -> bool:
        try:
            conn = await asyncio.wait_for(asyncpg.connect(dsn), timeout=3)
            await conn.close()
            return True
        except Exception:
            return False

    try:
        return asyncio.run(_probe())
    except RuntimeError:
        return False


REQUIRES_DB = pytest.mark.skipif(
    not _postgres_reachable(),
    reason="Postgres is not reachable from this host.",
)


@pytest_asyncio.fixture
async def db_session() -> AsyncIterator:
    from app.database.session import AsyncSessionLocal, init_engine

    if AsyncSessionLocal is None:
        init_engine()
    from app.database.session import AsyncSessionLocal as Maker

    async with Maker() as s:  # type: ignore[misc]
        yield s
        await s.rollback()


async def _seed_org(session, *, slug_hint: str):
    from datetime import datetime, timezone

    from app.database.models.organization import Organization, OrgPlan
    from app.database.models.organization_member import (
        MemberRole, MemberStatus, OrganizationMember,
    )
    from app.database.models.user import User

    user = User(
        cognito_sub=f"sub-{uuid.uuid4()}",
        email=f"{slug_hint}-{uuid.uuid4().hex[:6]}@x.com",
        full_name=slug_hint.title(),
    )
    session.add(user)
    await session.flush()
    org = Organization(
        name=f"{slug_hint.title()} Workspace",
        slug=f"{slug_hint}-{uuid.uuid4().hex[:8]}",
        plan=OrgPlan.free,
        owner_user_id=user.id,
    )
    session.add(org)
    await session.flush()
    session.add(
        OrganizationMember(
            organization_id=org.id,
            user_id=user.id,
            role=MemberRole.owner,
            status=MemberStatus.active,
            joined_at=datetime.now(timezone.utc),
        )
    )
    await session.flush()
    return user, org


@REQUIRES_DB
@pytest.mark.asyncio
async def test_status_transitions_pending_processing_processed(
    db_session, tmp_path, monkeypatch
):
    """Upload a TXT file → status walks pending → processing → processed
    and chunks materialise."""
    monkeypatch.delenv("S3_BUCKET", raising=False)
    monkeypatch.setenv("UPLOAD_DIR", str(tmp_path))

    from app.middleware.org_context import get_current_organization
    from server import app

    user, org = await _seed_org(db_session, slug_hint="p7")
    await db_session.commit()

    async def _override() -> OrgContext:
        return _ctx(org.id, user_id=user.id)

    app.dependency_overrides[get_current_organization] = _override
    try:
        with TestClient(app) as client:
            # Create a KB
            r = client.post("/api/knowledge-bases", json={"name": "KB", "status": "active"})
            assert r.status_code == 201
            kb_id = r.json()["id"]

            # Upload a sizeable TXT so we get >1 chunk
            text = ("OraOne is the unified AI agent platform. " * 50).encode()
            r = client.post(
                "/api/documents/upload",
                data={"knowledge_base_id": kb_id},
                files={"file": ("notes.txt", text, "text/plain")},
            )
            assert r.status_code == 201
            doc_id = r.json()["id"]

            # BackgroundTasks run *after* the response. Poll briefly.
            for _ in range(50):
                r = client.get(f"/api/documents/{doc_id}")
                if r.json()["status"] == "processed":
                    break
                await asyncio.sleep(0.05)

            doc = r.json()
            assert doc["status"] == "processed", doc
            assert doc["chunk_count"] > 0
            assert doc["processing_started_at"] is not None
            assert doc["processing_completed_at"] is not None
            assert doc["processing_time_ms"] is not None
            assert doc["processing_time_ms"] >= 0

            # Chunks endpoint returns them ordered
            r = client.get(f"/api/documents/{doc_id}/chunks")
            assert r.status_code == 200
            chunks = r.json()
            assert chunks and chunks[0]["chunk_index"] == 0
            # Required metadata keys
            md = chunks[0]["chunk_metadata"]
            assert "page" in md and "source_file" in md
    finally:
        app.dependency_overrides.pop(get_current_organization, None)


@REQUIRES_DB
@pytest.mark.asyncio
async def test_reprocess_endpoint_idempotent(db_session, tmp_path, monkeypatch):
    monkeypatch.delenv("S3_BUCKET", raising=False)
    monkeypatch.setenv("UPLOAD_DIR", str(tmp_path))

    from app.middleware.org_context import get_current_organization
    from server import app

    user, org = await _seed_org(db_session, slug_hint="reproc")
    await db_session.commit()

    async def _override() -> OrgContext:
        return _ctx(org.id, user_id=user.id)

    app.dependency_overrides[get_current_organization] = _override
    try:
        with TestClient(app) as client:
            kb_id = client.post(
                "/api/knowledge-bases", json={"name": "KB", "status": "active"}
            ).json()["id"]
            doc_id = client.post(
                "/api/documents/upload",
                data={"knowledge_base_id": kb_id},
                files={"file": ("a.txt", b"Hello there. " * 100, "text/plain")},
            ).json()["id"]

            for _ in range(50):
                if client.get(f"/api/documents/{doc_id}").json()["status"] == "processed":
                    break
                await asyncio.sleep(0.05)

            first_count = client.get(f"/api/documents/{doc_id}").json()["chunk_count"]
            assert first_count > 0

            # Manual reprocess
            r = client.post(f"/api/documents/{doc_id}/process")
            assert r.status_code == 202
            assert r.json()["status"] == "processing"

            for _ in range(50):
                if client.get(f"/api/documents/{doc_id}").json()["status"] == "processed":
                    break
                await asyncio.sleep(0.05)

            second_count = client.get(f"/api/documents/{doc_id}").json()["chunk_count"]
            # Reprocess should produce the same number of chunks for the same input
            assert second_count == first_count
    finally:
        app.dependency_overrides.pop(get_current_organization, None)
